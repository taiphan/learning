"""Anthropic brand styling for python-docx Word documents.

Matches anthropic_theme.py (PPTX): slate/cream palette, Poppins/Lora typography.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Brand colors (RGB)
SLATE = RGBColor(0x14, 0x14, 0x13)
CREAM = RGBColor(0xFA, 0xF9, 0xF5)
IVORY_SUBTLE = RGBColor(0xE8, 0xE6, 0xDC)
MID_GRAY = RGBColor(0xB0, 0xAE, 0xA5)
SLATE_MUTED = RGBColor(0x5E, 0x5D, 0x59)
CLAY = RGBColor(0xD9, 0x77, 0x57)
SKY = RGBColor(0x6A, 0x9B, 0xCC)
OLIVE = RGBColor(0x78, 0x8C, 0x5D)
FIG = RGBColor(0xC4, 0x66, 0x86)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_HEADING = "Poppins"
FONT_BODY = "Lora"
ACCENT_HEX = [CLAY, SKY, OLIVE, FIG]


def _run_font(run, *, name=FONT_BODY, size=11, bold=False, italic=False, color=SLATE):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def _para_bottom_border(paragraph, color_hex="E8E6DC", size=6):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _para_left_bar(paragraph, color_hex="D97757", width=24):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def setup_document(doc: Document) -> Document:
    """Apply Anthropic margins and base styles."""
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(0, 26), (1, 16), (2, 13), (3, 12)]:
        key = "Title" if level == 0 else f"Heading {level}"
        if key not in styles:
            continue
        h = styles[key]
        h.font.name = FONT_HEADING
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = SLATE
        h.paragraph_format.space_before = Pt(14 if level else 0)
        h.paragraph_format.space_after = Pt(10)
        h.paragraph_format.keep_with_next = True

    if "List Number" in styles:
        ln = styles["List Number"]
        ln.font.name = FONT_BODY
        ln.font.size = Pt(11)
        ln.font.color.rgb = SLATE
        ln.paragraph_format.space_after = Pt(4)

    return doc


def new_branded_document() -> Document:
    doc = Document()
    return setup_document(doc)


def add_cover(doc: Document, title: str, subtitle: str = "", meta: str = "") -> None:
    """Title block — Anthropic editorial cover."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    _run_font(run, name=FONT_HEADING, size=26, bold=True, color=SLATE)

    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sr = sp.add_run(subtitle)
        _run_font(sr, size=13, color=SLATE_MUTED)

    if meta:
        mp = doc.add_paragraph()
        mp.paragraph_format.space_before = Pt(2)
        mr = mp.add_run(meta)
        _run_font(mr, size=10, italic=True, color=MID_GRAY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(16)
    _para_bottom_border(rule, color_hex="D97757", size=18)


def add_section_heading(doc: Document, text: str, accent_index: int = 0) -> None:
    """Section header with clay/sky/olive left bar."""
    colors = ["D97757", "6A9BCC", "788C5D", "C46686"]
    hex_color = colors[accent_index % len(colors)]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    _para_left_bar(p, color_hex=hex_color, width=20)
    run = p.add_run(text)
    _run_font(run, name=FONT_HEADING, size=14, bold=True, color=SLATE)


def add_body(doc: Document, text: str, *, italic=False, muted=False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _run_font(run, size=11, italic=italic, color=SLATE_MUTED if muted else SLATE)


def add_field(doc: Document, label: str, placeholder: str = "") -> None:
    """Label + input line (bottom border UX)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    lr = p.add_run(f"{label} ")
    _run_font(lr, size=11, bold=True, color=SLATE)
    if placeholder:
        pr = p.add_run(placeholder)
        _run_font(pr, size=10, italic=True, color=SLATE_MUTED)
    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(10)
    if not placeholder:
        line.add_run("\u00a0")
    _para_bottom_border(line, color_hex="E8E6DC", size=8)


def add_numbered_items(doc: Document, items: list[str]) -> None:
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        nr = p.add_run(f"{i}. ")
        _run_font(nr, size=11, bold=True, color=CLAY)
        tr = p.add_run(item)
        _run_font(tr, size=11, color=SLATE)


def style_table(table, header_labels: list[str]) -> None:
    """Slate header row + cream body — Anthropic table UX."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, label in enumerate(header_labels):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(label)
        _run_font(run, name=FONT_HEADING, size=10, bold=True, color=WHITE)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "141413")
        hdr[i]._tc.get_or_add_tcPr().append(shading)

    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    _run_font(run, size=10, color=SLATE)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "FAF9F5")
            cell._tc.get_or_add_tcPr().append(shading)


def add_callout_box(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    _para_left_bar(p, color_hex="D97757", width=24)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F0EEE6")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    _run_font(run, size=10.5, color=SLATE)
